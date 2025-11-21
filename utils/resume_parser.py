import io
from typing import Union

from PyPDF2 import PdfReader
import docx


def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    text = []
    for page in reader.pages:
        text.append(page.extract_text() or "")
    return "\n".join(text)


def extract_text_from_docx(file) -> str:
    # file is a BytesIO from Streamlit uploader
    if isinstance(file, (io.BytesIO, io.BufferedReader)):
        doc = docx.Document(file)
    else:
        doc = docx.Document(io.BytesIO(file.read()))
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_resume(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        return ""
